"""
Supabase Vector Store for noteCHAT
Stores document chunks and their embeddings in Supabase PostgreSQL with pgvector
"""

import os
import logging
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
import numpy as np
import hashlib
from sentence_transformers import SentenceTransformer
from supabase import create_client, Client
from pathlib import Path
import json

# Document processing
import PyPDF2
from docx import Document as DocxDocument

# Text processing
import nltk
from nltk.tokenize import sent_tokenize
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Document chunk with metadata"""
    content: str
    source: str
    page_number: Optional[int] = None
    chunk_id: str = ""
    embedding: Optional[List[float]] = None

class SupabaseVectorStore:
    """Manages vector embeddings storage in Supabase"""
    
    def __init__(self, supabase_url: str, supabase_key: str):
        """
        Initialize Supabase vector store
        
        Args:
            supabase_url: Your Supabase project URL
            supabase_key: Your Supabase API key
        """
        self.supabase: Client = create_client(supabase_url, supabase_key)
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        logger.info("✅ Supabase Vector Store initialized")
        
    def setup_database(self):
        """
        Setup database tables and pgvector extension
        This creates the necessary table structure in Supabase
        """
        logger.info("🔧 Setting up database schema...")
        
        # SQL to create the table with pgvector
        create_table_sql = """
        -- Enable pgvector extension
        CREATE EXTENSION IF NOT EXISTS vector;
        
        -- Create document_chunks table
        CREATE TABLE IF NOT EXISTS document_chunks (
            id BIGSERIAL PRIMARY KEY,
            chunk_id TEXT UNIQUE NOT NULL,
            content TEXT NOT NULL,
            source TEXT NOT NULL,
            page_number INTEGER,
            embedding vector(384),  -- all-MiniLM-L6-v2 produces 384-dimensional vectors
            created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
            updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
        );
        
        -- Create index for faster similarity search
        CREATE INDEX IF NOT EXISTS document_chunks_embedding_idx 
        ON document_chunks 
        USING ivfflat (embedding vector_cosine_ops)
        WITH (lists = 100);
        
        -- Create index for source lookups
        CREATE INDEX IF NOT EXISTS document_chunks_source_idx 
        ON document_chunks(source);
        """
        
        logger.info("📋 SQL schema for Supabase:")
        logger.info(create_table_sql)
        logger.info("\n⚠️  Please run this SQL in your Supabase SQL Editor:")
        logger.info("1. Go to your Supabase Dashboard")
        logger.info("2. Navigate to SQL Editor")
        logger.info("3. Copy and paste the SQL above")
        logger.info("4. Click 'Run' to execute")
        
        return create_table_sql
    
    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        text_data = {"pages": [], "full_text": "", "metadata": {}}
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_data["metadata"]["total_pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            cleaned_text = self._clean_text(page_text)
                            if len(cleaned_text) > 50:
                                text_data["pages"].append({
                                    "page_number": page_num,
                                    "text": cleaned_text
                                })
                                text_data["full_text"] += f"\n--- Page {page_num} ---\n{cleaned_text}\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to extract page {page_num}: {e}")
                        
        except Exception as e:
            logger.error(f"❌ PDF extraction failed for {pdf_path}: {e}")
            
        return text_data
    
    def extract_text_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        text_data = {"paragraphs": [], "full_text": "", "metadata": {}}
        
        try:
            doc = DocxDocument(docx_path)
            text_data["metadata"]["total_paragraphs"] = len(doc.paragraphs)
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text and paragraph.text.strip():
                    cleaned_text = self._clean_text(paragraph.text)
                    if len(cleaned_text) > 10:
                        text_data["paragraphs"].append({
                            "paragraph_number": para_num,
                            "text": cleaned_text
                        })
                        text_data["full_text"] += f"{cleaned_text}\n"
                    
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed for {docx_path}: {e}")
            
        return text_data
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text.strip())
        # Remove problematic characters
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        # Final cleanup
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def create_chunks(self, text_data: Dict[str, Any], source: str) -> List[DocumentChunk]:
        """Create semantic chunks from text data"""
        chunks = []
        
        if "pages" in text_data:  # PDF
            for page_info in text_data["pages"]:
                page_chunks = self._chunk_text(
                    page_info["text"], 
                    source, 
                    page_info["page_number"]
                )
                chunks.extend(page_chunks)
                
        elif "paragraphs" in text_data:  # DOCX
            full_text = "\n".join([p["text"] for p in text_data["paragraphs"]])
            chunks = self._chunk_text(full_text, source)
                
        return chunks
    
    def _chunk_text(self, text: str, source: str, page_number: Optional[int] = None) -> List[DocumentChunk]:
        """Split text into chunks with overlap"""
        try:
            sentences = sent_tokenize(text)
        except:
            sentences = text.split('. ')
            
        chunks = []
        target_length = 200  # words per chunk
        overlap_sentences = 1
        
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence.split())
            
            if current_length + sentence_length > target_length and current_chunk:
                chunk_text = " ".join(current_chunk)
                if len(chunk_text.strip()) > 50:
                    # Generate unique chunk_id using hash of content
                    content_hash = hashlib.md5(chunk_text.encode()).hexdigest()[:8]
                    chunk = DocumentChunk(
                        content=chunk_text,
                        source=source,
                        page_number=page_number,
                        chunk_id=f"{source}_{page_number or 0}_{content_hash}"
                    )
                    chunks.append(chunk)
                
                if len(current_chunk) > overlap_sentences:
                    current_chunk = current_chunk[-overlap_sentences:]
                    current_length = sum(len(s.split()) for s in current_chunk)
                else:
                    current_chunk = []
                    current_length = 0
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            if len(chunk_text.strip()) > 50:
                # Generate unique chunk_id using hash of content
                content_hash = hashlib.md5(chunk_text.encode()).hexdigest()[:8]
                chunk = DocumentChunk(
                    content=chunk_text,
                    source=source,
                    page_number=page_number,
                    chunk_id=f"{source}_{page_number or 0}_{content_hash}"
                )
                chunks.append(chunk)
            
        return chunks
    
    def generate_embeddings(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Generate embeddings for chunks"""
        logger.info(f"🧠 Generating embeddings for {len(chunks)} chunks...")
        
        texts = [chunk.content for chunk in chunks]
        embeddings = self.embedding_model.encode(texts, show_progress_bar=True)
        
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding.tolist()
        
        logger.info("✅ Embeddings generated successfully")
        return chunks
    
    def insert_chunks(self, chunks: List[DocumentChunk]) -> bool:
        """Insert chunks with embeddings into Supabase"""
        try:
            logger.info(f"📤 Uploading {len(chunks)} chunks to Supabase...")
            
            # Prepare data for insertion
            records = []
            for chunk in chunks:
                record = {
                    "chunk_id": chunk.chunk_id,
                    "content": chunk.content,
                    "source": chunk.source,
                    "page_number": chunk.page_number,
                    "embedding": chunk.embedding
                }
                records.append(record)
            
            # Insert in batches with upsert to handle duplicates
            batch_size = 100
            successful = 0
            skipped = 0
            
            for i in range(0, len(records), batch_size):
                batch = records[i:i + batch_size]
                try:
                    # Use upsert to update existing or insert new
                    result = self.supabase.table("document_chunks").upsert(
                        batch,
                        on_conflict="chunk_id"
                    ).execute()
                    successful += len(batch)
                    logger.info(f"✅ Uploaded batch {i//batch_size + 1}/{(len(records)-1)//batch_size + 1}")
                except Exception as batch_error:
                    logger.warning(f"⚠️ Batch {i//batch_size + 1} had issues: {batch_error}")
                    skipped += len(batch)
            
            logger.info(f"✅ Successfully processed {successful} chunks (skipped {skipped} duplicates)")
            return True
            
        except Exception as e:
            logger.error(f"❌ Failed to insert chunks: {e}")
            return False
    
    def process_documents_folder(self, folder_path: str) -> bool:
        """Process all documents in a folder and store in Supabase"""
        folder = Path(folder_path)
        
        if not folder.exists():
            logger.error(f"❌ Folder not found: {folder_path}")
            return False
        
        logger.info(f"📁 Processing documents from: {folder_path}")
        
        all_chunks = []
        file_count = 0
        
        # Process all PDF and DOCX files
        for file_path in folder.glob("*"):
            if file_path.suffix.lower() in ['.pdf', '.docx']:
                logger.info(f"📄 Processing: {file_path.name}")
                file_count += 1
                
                try:
                    # Extract text
                    if file_path.suffix.lower() == '.pdf':
                        text_data = self.extract_text_from_pdf(str(file_path))
                    else:
                        text_data = self.extract_text_from_docx(str(file_path))
                    
                    # Create chunks
                    if text_data.get("full_text"):
                        chunks = self.create_chunks(text_data, file_path.name)
                        all_chunks.extend(chunks)
                        logger.info(f"✅ Created {len(chunks)} chunks from {file_path.name}")
                    
                except Exception as e:
                    logger.error(f"❌ Failed to process {file_path.name}: {e}")
        
        if not all_chunks:
            logger.warning("⚠️ No chunks created from documents")
            return False
        
        logger.info(f"📊 Total: {len(all_chunks)} chunks from {file_count} files")
        
        # Generate embeddings
        all_chunks = self.generate_embeddings(all_chunks)
        
        # Upload to Supabase
        return self.insert_chunks(all_chunks)
    
    def search_similar(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for similar chunks using vector similarity"""
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Search in Supabase using RPC function
            result = self.supabase.rpc(
                "match_documents",
                {
                    "query_embedding": query_embedding.tolist(),
                    "match_count": top_k
                }
            ).execute()
            
            return result.data
            
        except Exception as e:
            logger.error(f"❌ Search failed: {e}")
            return []
    
    def clear_all_chunks(self) -> bool:
        """Clear all chunks from database (use with caution!)"""
        try:
            logger.warning("⚠️ Clearing all chunks from database...")
            self.supabase.table("document_chunks").delete().neq("id", 0).execute()
            logger.info("✅ Database cleared")
            return True
        except Exception as e:
            logger.error(f"❌ Failed to clear database: {e}")
            return False


if __name__ == "__main__":
    # Example usage
    print("🚀 Supabase Vector Store for noteCHAT")
    print("\n📋 Setup Instructions:")
    print("1. Create a new Supabase project at https://supabase.com")
    print("2. Get your project URL and API key from Settings > API")
    print("3. Set environment variables:")
    print("   export SUPABASE_URL='your-project-url'")
    print("   export SUPABASE_KEY='your-anon-key'")
    print("\n4. Run the setup to create database schema")
    print("5. Process your documents")
