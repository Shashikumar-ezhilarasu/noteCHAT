"""
Robust RAG Pipeline with Enhanced Document Processing
Provides high-quality answers with proper references and semantic understanding
"""

import os
import shutil
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import re
import json
from dataclasses import dataclass, asdict
from collections import defaultdict

# Document processing
import PyPDF2
import docx
from docx import Document

# Text processing
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

# ML and similarity
import numpy as np
from sentence_transformers import SentenceTransformer, CrossEncoder
from sklearn.metrics.pairwise import cosine_similarity

@dataclass
class DocumentChunk:
    """Enhanced document chunk with metadata"""
    content: str
    source: str
    page_number: Optional[int] = None
    section: Optional[str] = None
    chunk_id: str = ""
    confidence_score: float = 0.0

class RobustRAGPipeline:
    """Robust RAG pipeline with enhanced document processing"""
    
    def __init__(self):
        self.logger = self._setup_logging()
        self.chunks: List[DocumentChunk] = []
        self.embedding_model: Optional[SentenceTransformer] = None
        self.cross_encoder_model: Optional[CrossEncoder] = None
        self.embedding_matrix = None
        self.lemmatizer = WordNetLemmatizer()
        try:
            self.stop_words = set(stopwords.words('english'))
        except:
            self.stop_words = set()
        
        self.cache_dir = Path("cache")
        self.chunks_cache_path = self.cache_dir / "chunks.json"
        self.embeddings_cache_path = self.cache_dir / "embeddings.npy"
        
    def _setup_logging(self) -> logging.Logger:
        """Setup enhanced logging"""
        logging.basicConfig(level=logging.INFO)
        logger = logging.getLogger("robust_rag")
        return logger
        
    def download_documents_from_local(self) -> List[str]:
        """Download documents from local directory"""
        local_files = []
        try:
            self.logger.info("🔄 Using local files")
            notes_dir = Path("../NOTES")
            if notes_dir.exists():
                os.makedirs("downloads", exist_ok=True)
                for file_path in notes_dir.glob("*"):
                    if file_path.suffix.lower() in ['.pdf', '.docx']:
                        dest_path = f"downloads/{file_path.name}"
                        shutil.copy2(file_path, dest_path)
                        local_files.append(dest_path)
                        self.logger.info(f"📁 Copied local: {file_path.name}")
            else:
                self.logger.warning("❌ NOTES directory not found")
                
        except Exception as e:
            self.logger.error(f"❌ Document download failed: {e}")
            
        return local_files
        
    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Enhanced PDF text extraction"""
        text_data = {"pages": [], "full_text": "", "metadata": {}}
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_data["metadata"]["total_pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            # Clean and format text
                            cleaned_text = self._clean_text(page_text)
                            if len(cleaned_text) > 50:  # Only keep substantial content
                                text_data["pages"].append({
                                    "page_number": page_num,
                                    "text": cleaned_text
                                })
                                text_data["full_text"] += f"\n--- Page {page_num} ---\n{cleaned_text}\n"
                    except Exception as e:
                        self.logger.warning(f"⚠️ Failed to extract page {page_num}: {e}")
                        
        except Exception as e:
            self.logger.error(f"❌ PDF extraction failed for {pdf_path}: {e}")
            
        return text_data
        
    def extract_text_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """Enhanced DOCX text extraction"""
        text_data = {"paragraphs": [], "full_text": "", "metadata": {}}
        
        try:
            doc = Document(docx_path)
            text_data["metadata"]["total_paragraphs"] = len(doc.paragraphs)
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text and paragraph.text.strip():
                    cleaned_text = self._clean_text(paragraph.text)
                    if len(cleaned_text) > 10:  # Only keep meaningful content
                        text_data["paragraphs"].append({
                            "paragraph_number": para_num,
                            "text": cleaned_text
                        })
                        text_data["full_text"] += f"{cleaned_text}\n"
                    
        except Exception as e:
            self.logger.error(f"❌ DOCX extraction failed for {docx_path}: {e}")
            
        return text_data
        
    def _fix_pdf_spacing(self, text: str) -> str:
        """Specialized function to fix PDF text extraction spacing issues"""
        # First pass: identify and fix obvious concatenations
        
        # Fix common patterns where words are stuck together
        # Pattern: lowercase letter followed by uppercase (camelCase)
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Pattern: letter followed by number or number followed by letter
        text = re.sub(r'(\d)([A-Za-z])', r'\1 \2', text)
        text = re.sub(r'([A-Za-z])(\d)', r'\1 \2', text)
        
        # Fix words stuck together by looking for common English patterns
        # Insert space before common word endings when preceded by a letter
        word_endings = ['ing', 'tion', 'sion', 'ment', 'ness', 'able', 'ible', 'ful', 'less', 'ous', 'ive', 'ed', 'er', 'est', 'ly', 'al', 'ic', 'ary', 'ory']
        for ending in word_endings:
            pattern = f'([a-z])({ending})([A-Z][a-z])'
            text = re.sub(pattern, rf'\1\2 \3', text)
        
        # Fix concatenated common words
        common_patterns = [
            (r'([a-z])(the)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(and)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(that)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(with)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(this)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(which)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(when)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(where)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(what)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(how)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(can)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(will)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(are)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(is)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(be)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(have)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(has)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(in)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(on)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(of)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(to)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(for)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(by)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(from)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(using)([A-Z])', r'\1 \2 \3'),
        ]
        
        for pattern, replacement in common_patterns:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        # Fix punctuation spacing
        text = re.sub(r'\.([A-Z])', r'. \1', text)
        text = re.sub(r',([A-Za-z])', r', \1', text)
        text = re.sub(r':([A-Za-z])', r': \1', text)
        text = re.sub(r';([A-Za-z])', r'; \1', text)
        text = re.sub(r'\!([A-Za-z])', r'! \1', text)
        text = re.sub(r'\?([A-Za-z])', r'? \1', text)
        
        # Handle parentheses spacing
        text = re.sub(r'\)([A-Za-z])', r') \1', text)
        text = re.sub(r'([A-Za-z])\(', r'\1 (', text)
        
        return text

    def _clean_text(self, text: str) -> str:
        """Advanced text cleaning and formatting with PDF-specific fixes"""
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Apply PDF-specific spacing fixes
        text = self._fix_pdf_spacing(text)
        
        # Remove problematic characters but preserve structure
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        
        # Final cleanup
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
        
    def create_semantic_chunks(self, text_data: Dict[str, Any], source: str) -> List[DocumentChunk]:
        """Create semantic chunks with better overlap"""
        chunks = []
        
        if "pages" in text_data:  # PDF
            for page_info in text_data["pages"]:
                page_chunks = self._chunk_text_intelligently(
                    page_info["text"], 
                    source, 
                    page_info["page_number"]
                )
                chunks.extend(page_chunks)
                
        elif "paragraphs" in text_data:  # DOCX
            full_text = "\n".join([p["text"] for p in text_data["paragraphs"]])
            chunks = self._chunk_text_intelligently(full_text, source)
                
        return chunks
        
    def _chunk_text_intelligently(self, text: str, source: str, page_number: Optional[int] = None) -> List[DocumentChunk]:
        """Create intelligent chunks with proper overlap"""
        try:
            sentences = sent_tokenize(text)
        except:
            # Fallback if NLTK fails
            sentences = text.split('. ')
            
        chunks = []
        
        # Parameters for chunking
        target_length = 200  # Target words per chunk
        overlap_sentences = 1  # Number of sentences to overlap
        
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence.split())
            
            # Check if adding this sentence would exceed target length
            if current_length + sentence_length > target_length and current_chunk:
                # Create chunk
                chunk_text = " ".join(current_chunk)
                if len(chunk_text.strip()) > 50:  # Only keep substantial chunks
                    chunk = DocumentChunk(
                        content=chunk_text,
                        source=source,
                        page_number=page_number,
                        chunk_id=f"{source}_{len(chunks)}"
                    )
                    chunks.append(chunk)
                
                # Keep overlap for context
                if len(current_chunk) > overlap_sentences:
                    current_chunk = current_chunk[-overlap_sentences:]
                    current_length = sum(len(s.split()) for s in current_chunk)
                else:
                    current_chunk = []
                    current_length = 0
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        # Final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            if len(chunk_text.strip()) > 50:
                chunk = DocumentChunk(
                    content=chunk_text,
                    source=source,
                    page_number=page_number,
                    chunk_id=f"{source}_{len(chunks)}"
                )
                chunks.append(chunk)
            
        return chunks
        
    def process_documents(self, file_paths: List[str]) -> bool:
        """Process all documents with enhanced extraction"""
        self.logger.info(f"🔄 Processing {len(file_paths)} documents...")
        
        total_chunks = 0
        processed_docs = 0
        
        for file_path in file_paths:
            try:
                self.logger.info(f"📄 Processing: {os.path.basename(file_path)}")
                
                if file_path.lower().endswith('.pdf'):
                    text_data = self.extract_text_from_pdf(file_path)
                elif file_path.lower().endswith('.docx'):
                    text_data = self.extract_text_from_docx(file_path)
                else:
                    self.logger.warning(f"⚠️ Unsupported file type: {file_path}")
                    continue
                
                if text_data["full_text"].strip():
                    doc_chunks = self.create_semantic_chunks(text_data, os.path.basename(file_path))
                    if doc_chunks:
                        self.chunks.extend(doc_chunks)
                        total_chunks += len(doc_chunks)
                        processed_docs += 1
                        self.logger.info(f"✅ Processed {os.path.basename(file_path)}: {len(doc_chunks)} chunks")
                    else:
                        self.logger.warning(f"⚠️ No chunks created from {file_path}")
                else:
                    self.logger.warning(f"⚠️ No text extracted from {file_path}")
                    
            except Exception as e:
                self.logger.error(f"❌ Failed to process {file_path}: {e}")
        
        self.logger.info(f"📚 Total: {total_chunks} chunks from {processed_docs} documents")
        
        # Log some sample chunks for debugging
        if self.chunks:
            self.logger.info(f"📝 Sample chunk: {self.chunks[0].content[:100]}...")
            
        return len(self.chunks) > 0
        
    def create_search_index(self) -> bool:
        """Create search index with sentence transformer embeddings"""
        try:
            if not self.chunks:
                self.logger.error("❌ No chunks to index")
                return False
            
            self.logger.info("🧠 Loading sentence transformer model...")
            # Using a lightweight but powerful model
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            
            self.logger.info("🧠 Loading Cross-Encoder model for re-ranking...")
            self.cross_encoder_model = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')

            self.logger.info("🧠 Creating embeddings for all chunks...")
            # Prepare texts for embedding
            texts = [chunk.content for chunk in self.chunks]
            self.embedding_matrix = self.embedding_model.encode(texts, show_progress_bar=True)
            
            self.logger.info(f"✅ Search index created with embeddings: {self.embedding_matrix.shape}")
            
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Failed to create search index: {e}")
            return False
            
    def find_relevant_chunks(self, query: str, top_k: int = 5) -> List[Tuple[DocumentChunk, float]]:
        """Find most relevant chunks using semantic search and re-ranking"""
        try:
            if self.embedding_model is None or self.embedding_matrix is None or self.cross_encoder_model is None:
                self.logger.error("❌ Search index or cross-encoder not available")
                return []
            
            self.logger.info(f"🔍 Stage 1: Initial search for: {query}")
            
            # Stage 1: Fast retrieval with sentence transformer
            query_embedding = self.embedding_model.encode([query])
            similarities = cosine_similarity(query_embedding, self.embedding_matrix).flatten()
            
            # Get a larger pool of candidates for re-ranking
            candidate_indices = similarities.argsort()[-top_k * 4:][::-1]
            
            # Stage 2: Re-ranking with Cross-Encoder
            self.logger.info(f"🔍 Stage 2: Re-ranking {len(candidate_indices)} candidates with Cross-Encoder...")
            
            cross_encoder_inputs = [[query, self.chunks[idx].content] for idx in candidate_indices]
            cross_encoder_scores = self.cross_encoder_model.predict(cross_encoder_inputs)

            # Apply sigmoid to convert scores to a 0-1 confidence range
            confidence_scores = 1 / (1 + np.exp(-cross_encoder_scores))
            
            # Combine candidates with their new scores
            reranked_results = list(zip(candidate_indices, confidence_scores))
            
            # Sort by the new confidence scores in descending order
            reranked_results.sort(key=lambda x: x[1], reverse=True)
            
            # Log re-ranking stats
            max_rerank_score = reranked_results[0][1] if reranked_results else 0
            self.logger.info(f"📊 Re-ranking stats - Max Score: {max_rerank_score:.3f}")

            # Get top matches after re-ranking
            results = []
            for idx, score in reranked_results[:top_k]:
                chunk = self.chunks[idx]
                chunk.confidence_score = float(score)
                results.append((chunk, float(score)))
                self.logger.info(f"📋 Final Match {len(results)}: {score:.3f} - {chunk.content[:50]}...")
            
            return results
            
        except Exception as e:
            self.logger.error(f"❌ Search failed: {e}")
            return []
            
    def generate_comprehensive_answer(self, query: str, relevant_chunks: List[Tuple[DocumentChunk, float]]) -> Dict[str, Any]:
        """Generate comprehensive answer with proper references"""
        if not relevant_chunks:
            return {
                "answer": "I don't have enough information to answer this question based on the available documents.",
                "sources": [],
                "confidence": 0.0,
                "references": []
            }
        
        # Build comprehensive answer
        answer_parts = []
        references = []
        sources = set()
        
        for i, (chunk, score) in enumerate(relevant_chunks, 1):
            # Format content with source reference
            reference_info = f"**Source {i}** (from {chunk.source}"
            if chunk.page_number:
                reference_info += f", Page {chunk.page_number}"
            reference_info += f", Confidence: {score:.2f})"
            
            answer_parts.append(f"{reference_info}:\n{chunk.content}")
            
            references.append({
                "source": chunk.source,
                "page": chunk.page_number,
                "confidence": score,
                "rank": i
            })
            sources.add(chunk.source)
        
        # Calculate overall confidence
        avg_confidence = sum(score for _, score in relevant_chunks) / len(relevant_chunks)
        
        answer = "\n\n".join(answer_parts)
        
        return {
            "answer": answer,
            "sources": list(sources),
            "confidence": avg_confidence,
            "references": references
        }
        
    def query(self, question: str) -> Dict[str, Any]:
        """Process query and return comprehensive answer"""
        try:
            self.logger.info(f"🔍 Processing query: {question}")
            
            # Find relevant chunks
            relevant_chunks = self.find_relevant_chunks(question, top_k=3)
            
            if not relevant_chunks:
                self.logger.warning("❌ No relevant chunks found")
                return {
                    "answer": "I couldn't find relevant information to answer this question in the available documents.",
                    "sources": [],
                    "confidence": 0.0
                }
            
            # Generate comprehensive answer
            result = self.generate_comprehensive_answer(question, relevant_chunks)
            
            self.logger.info(f"✅ Answer generated with {len(result['sources'])} sources, confidence: {result['confidence']:.3f}")
            return result
            
        except Exception as e:
            self.logger.error(f"❌ Query processing failed: {e}")
            return {
                "answer": "An error occurred while processing your question.",
                "sources": [],
                "confidence": 0.0
            }

    def _save_to_cache(self):
        """Save the processed chunks and embeddings to disk."""
        try:
            self.logger.info(f"💾 Saving knowledge base to cache directory: {self.cache_dir}")
            self.cache_dir.mkdir(exist_ok=True)

            # Save chunks as JSON
            with open(self.chunks_cache_path, 'w', encoding='utf-8') as f:
                json.dump([asdict(chunk) for chunk in self.chunks], f, indent=4)

            # Save embeddings as a numpy file
            if self.embedding_matrix is not None:
                np.save(self.embeddings_cache_path, self.embedding_matrix)
            
            self.logger.info("✅ Knowledge base saved successfully.")
        except Exception as e:
            self.logger.error(f"❌ Failed to save knowledge base to cache: {e}")

    def _load_from_cache(self) -> bool:
        """Load the processed chunks and embeddings from disk."""
        try:
            self.logger.info("🔄 Loading knowledge base from cache...")

            # Load chunks from JSON
            with open(self.chunks_cache_path, 'r', encoding='utf-8') as f:
                chunks_data = json.load(f)
                self.chunks = [DocumentChunk(**data) for data in chunks_data]

            # Load embeddings
            self.embedding_matrix = np.load(self.embeddings_cache_path)
            
            # Load the sentence transformer model
            self.logger.info("🧠 Loading sentence transformer model...")
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

            self.logger.info("🧠 Loading Cross-Encoder model for re-ranking...")
            self.cross_encoder_model = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')

            self.logger.info(f"✅ Knowledge base loaded successfully from cache. {len(self.chunks)} chunks.")
            return True
        except Exception as e:
            self.logger.error(f"❌ Failed to load knowledge base from cache: {e}")
            return False
            
    def initialize(self) -> bool:
        """Initialize the robust RAG pipeline"""
        try:
            self.logger.info("🚀 Initializing Robust RAG Pipeline...")
            
            # Check if cache exists and load from it
            if self.chunks_cache_path.exists() and self.embeddings_cache_path.exists():
                if self._load_from_cache():
                    self.logger.info("🎉 Robust RAG Pipeline initialized successfully from cache!")
                    return True
                else:
                    self.logger.warning("⚠️ Failed to load from cache, rebuilding...")

            # If cache doesn't exist or loading failed, build it
            self.logger.info("🛠️ Knowledge base not found in cache. Building from scratch...")
            
            # Download documents
            file_paths = self.download_documents_from_local()
            if not file_paths:
                self.logger.error("❌ No documents found")
                return False
            
            # Process documents
            if not self.process_documents(file_paths):
                self.logger.error("❌ Document processing failed")
                return False
            
            # Create search index
            if not self.create_search_index():
                self.logger.error("❌ Search index creation failed")
                return False
            
            # Load the cross-encoder model after building the index
            self.logger.info("🧠 Loading Cross-Encoder model for re-ranking...")
            self.cross_encoder_model = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')

            # Save the newly built knowledge base to cache
            self._save_to_cache()
            
            self.logger.info("🎉 Robust RAG Pipeline initialized successfully!")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Initialization failed: {e}")
            return False
