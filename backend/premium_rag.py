"""
Premium RAG Pipeline with Advanced Document Processing
Provides high-quality answers with proper references and semantic understanding
"""

import os
import shutil
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import re
import json
from dataclasses import dataclass
from collections import defaultdict

# Document processing
import PyPDF2
import docx
from docx import Document

# Text processing
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

# ML and similarity
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.cluster import KMeans

# Firebase
import firebase_admin
from firebase_admin import credentials, storage

@dataclass
class DocumentChunk:
    """Enhanced document chunk with metadata"""
    content: str
    source: str
    page_number: Optional[int] = None
    section: Optional[str] = None
    chunk_id: str = ""
    confidence_score: float = 0.0
    keywords: List[str] = None
    
    def __post_init__(self):
        if self.keywords is None:
            self.keywords = []

class PremiumRAGPipeline:
    """Premium RAG pipeline with advanced document processing and semantic understanding"""
    
    def __init__(self):
        self.logger = self._setup_logging()
        self.chunks: List[DocumentChunk] = []
        self.vectorizer: Optional[TfidfVectorizer] = None
        self.tfidf_matrix = None
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
        self.firebase_initialized = False
        
        # ML Topics and Keywords
        self.ml_topics = {
            'supervised_learning': ['supervised', 'classification', 'regression', 'labeled', 'training', 'target'],
            'unsupervised_learning': ['unsupervised', 'clustering', 'dimensionality', 'unlabeled', 'pattern'],
            'kmeans': ['k-means', 'kmeans', 'centroid', 'cluster', 'elbow', 'wcss'],
            'hierarchical': ['hierarchical', 'dendrogram', 'agglomerative', 'divisive', 'linkage'],
            'hmm': ['hidden markov', 'hmm', 'states', 'transition', 'emission', 'viterbi'],
            'neural_networks': ['neural', 'network', 'deep', 'learning', 'perceptron', 'backpropagation'],
            'statistics': ['mean', 'median', 'variance', 'distribution', 'probability', 'quantile']
        }
        
    def _setup_logging(self) -> logging.Logger:
        """Setup enhanced logging"""
        logging.basicConfig(level=logging.INFO)
        logger = logging.getLogger("premium_rag")
        return logger
        
    def initialize_firebase(self) -> bool:
        """Initialize Firebase with error handling"""
        try:
            if not firebase_admin._apps:
                cred_path = "firebase-config.json"
                if os.path.exists(cred_path):
                    cred = credentials.Certificate(cred_path)
                    firebase_admin.initialize_app(cred, {
                        'storageBucket': 'notechat-26c38.firebasestorage.app'
                    })
                    self.firebase_initialized = True
                    self.logger.info("✅ Firebase initialized successfully")
                    return True
        except Exception as e:
            self.logger.error(f"❌ Firebase initialization failed: {e}")
        return False
        
    def download_documents_from_firebase(self) -> List[str]:
        """Download documents from Firebase Storage"""
        local_files = []
        try:
            if self.firebase_initialized:
                bucket = storage.bucket()
                blobs = bucket.list_blobs()
                
                os.makedirs("downloads", exist_ok=True)
                
                for blob in blobs:
                    if blob.name.endswith(('.pdf', '.docx')):
                        local_path = f"downloads/{blob.name}"
                        blob.download_to_filename(local_path)
                        local_files.append(local_path)
                        self.logger.info(f"📥 Downloaded: {blob.name}")
            else:
                self.logger.warning("🔄 Using local files")
                notes_dir = Path("../NOTES")
                if notes_dir.exists():
                    os.makedirs("downloads", exist_ok=True)
                    for file_path in notes_dir.glob("*"):
                        if file_path.suffix.lower() in ['.pdf', '.docx']:
                            dest_path = f"downloads/{file_path.name}"
                            shutil.copy2(file_path, dest_path)
                            local_files.append(dest_path)
                            self.logger.info(f"📁 Copied local: {file_path.name}")
                            
        except Exception as e:
            self.logger.error(f"❌ Document download failed: {e}")
            
        return local_files
        
    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Enhanced PDF text extraction with page tracking"""
        text_data = {"pages": [], "full_text": "", "metadata": {}}
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_data["metadata"]["total_pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            # Clean and format text
                            cleaned_text = self._clean_text(page_text)
                            text_data["pages"].append({
                                "page_number": page_num,
                                "text": cleaned_text
                            })
                            text_data["full_text"] += f"\n--- Page {page_num} ---\n{cleaned_text}\n"
                    except Exception as e:
                        self.logger.warning(f"⚠️ Failed to extract page {page_num}: {e}")
                        
        except Exception as e:
            self.logger.error(f"❌ PDF extraction failed for {pdf_path}: {e}")
            
        return text_data
        
    def extract_text_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """Enhanced DOCX text extraction with structure preservation"""
        text_data = {"paragraphs": [], "full_text": "", "metadata": {}}
        
        try:
            doc = Document(docx_path)
            text_data["metadata"]["total_paragraphs"] = len(doc.paragraphs)
            
            current_section = ""
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    cleaned_text = self._clean_text(paragraph.text)
                    
                    # Detect sections/headings
                    if self._is_heading(paragraph):
                        current_section = cleaned_text
                        
                    text_data["paragraphs"].append({
                        "paragraph_number": para_num,
                        "text": cleaned_text,
                        "section": current_section
                    })
                    text_data["full_text"] += f"{cleaned_text}\n"
                    
        except Exception as e:
            self.logger.error(f"❌ DOCX extraction failed for {docx_path}: {e}")
            
        return text_data
        
    def _clean_text(self, text: str) -> str:
        """Advanced text cleaning and formatting"""
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Fix common OCR issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Split concatenated words
        text = re.sub(r'(\d+)([A-Za-z])', r'\1 \2', text)  # Space between numbers and letters
        text = re.sub(r'([A-Za-z])(\d+)', r'\1 \2', text)  # Space between letters and numbers
        
        # Clean special characters but preserve structure
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
        
    def _is_heading(self, paragraph) -> bool:
        """Detect if a paragraph is a heading"""
        text = paragraph.text.strip()
        
        # Check for heading indicators
        if len(text) < 100 and any([
            text.isupper(),
            text.endswith(':'),
            paragraph.style.name.startswith('Heading'),
            re.match(r'^\d+\.', text),  # Numbered sections
            re.match(r'^[A-Z][a-z]+ \d+', text)  # "Unit 1", "Chapter 2"
        ]):
            return True
        return False
        
    def create_intelligent_chunks(self, text_data: Dict[str, Any], source: str) -> List[DocumentChunk]:
        """Create intelligent chunks with context preservation"""
        chunks = []
        
        if "pages" in text_data:  # PDF
            for page_info in text_data["pages"]:
                page_chunks = self._chunk_text_semantically(
                    page_info["text"], 
                    source, 
                    page_info["page_number"]
                )
                chunks.extend(page_chunks)
                
        elif "paragraphs" in text_data:  # DOCX
            current_section = ""
            section_content = []
            
            for para_info in text_data["paragraphs"]:
                if self._is_heading_text(para_info["text"]):
                    # Process previous section
                    if section_content:
                        section_text = " ".join(section_content)
                        section_chunks = self._chunk_text_semantically(
                            section_text, source, section=current_section
                        )
                        chunks.extend(section_chunks)
                    
                    # Start new section
                    current_section = para_info["text"]
                    section_content = []
                else:
                    section_content.append(para_info["text"])
            
            # Process final section
            if section_content:
                section_text = " ".join(section_content)
                section_chunks = self._chunk_text_semantically(
                    section_text, source, section=current_section
                )
                chunks.extend(section_chunks)
                
        return chunks
        
    def _is_heading_text(self, text: str) -> bool:
        """Check if text appears to be a heading"""
        return len(text) < 100 and any([
            text.isupper(),
            text.endswith(':'),
            re.match(r'^\d+\.', text),
            re.match(r'^[A-Z][a-z]+ \d+', text)
        ])
        
    def _chunk_text_semantically(self, text: str, source: str, page_number: Optional[int] = None, section: Optional[str] = None) -> List[DocumentChunk]:
        """Create semantic chunks with overlapping context"""
        sentences = sent_tokenize(text)
        chunks = []
        
        # Group sentences into meaningful chunks
        current_chunk = []
        current_length = 0
        target_length = 300  # Target words per chunk
        overlap_length = 50   # Overlap between chunks
        
        for sentence in sentences:
            sentence_length = len(sentence.split())
            
            if current_length + sentence_length > target_length and current_chunk:
                # Create chunk
                chunk_text = " ".join(current_chunk)
                chunk = DocumentChunk(
                    content=chunk_text,
                    source=source,
                    page_number=page_number,
                    section=section,
                    chunk_id=f"{source}_{len(chunks)}",
                    keywords=self._extract_keywords(chunk_text)
                )
                chunks.append(chunk)
                
                # Keep overlap for context
                overlap_sentences = current_chunk[-2:] if len(current_chunk) > 2 else current_chunk
                current_chunk = overlap_sentences + [sentence]
                current_length = sum(len(s.split()) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_length += sentence_length
        
        # Final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk = DocumentChunk(
                content=chunk_text,
                source=source,
                page_number=page_number,
                section=section,
                chunk_id=f"{source}_{len(chunks)}",
                keywords=self._extract_keywords(chunk_text)
            )
            chunks.append(chunk)
            
        return chunks
        
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract relevant keywords from text"""
        words = word_tokenize(text.lower())
        keywords = []
        
        for word in words:
            if (len(word) > 3 and 
                word.isalpha() and 
                word not in self.stop_words):
                lemmatized = self.lemmatizer.lemmatize(word)
                keywords.append(lemmatized)
        
        # Add topic-specific keywords
        for topic, topic_keywords in self.ml_topics.items():
            for keyword in topic_keywords:
                if keyword in text.lower():
                    keywords.append(topic)
                    break
        
        return list(set(keywords))
        
    def process_documents(self, file_paths: List[str]) -> bool:
        """Process all documents with enhanced extraction"""
        self.logger.info(f"🔄 Processing {len(file_paths)} documents...")
        
        total_chunks = 0
        processed_docs = 0
        
        for file_path in file_paths:
            try:
                self.logger.info(f"📄 Processing: {os.path.basename(file_path)}")
                
                if file_path.lower().endswith('.pdf'):
                    text_data = self.extract_text_from_pdf(file_path)
                elif file_path.lower().endswith('.docx'):
                    text_data = self.extract_text_from_docx(file_path)
                else:
                    self.logger.warning(f"⚠️ Unsupported file type: {file_path}")
                    continue
                
                if text_data["full_text"].strip():
                    doc_chunks = self.create_intelligent_chunks(text_data, os.path.basename(file_path))
                    self.chunks.extend(doc_chunks)
                    total_chunks += len(doc_chunks)
                    processed_docs += 1
                    self.logger.info(f"✅ Processed {os.path.basename(file_path)}: {len(doc_chunks)} chunks")
                else:
                    self.logger.warning(f"⚠️ No text extracted from {file_path}")
                    
            except Exception as e:
                self.logger.error(f"❌ Failed to process {file_path}: {e}")
        
        self.logger.info(f"📚 Total: {total_chunks} chunks from {processed_docs} documents")
        return len(self.chunks) > 0
        
    def create_search_index(self) -> bool:
        """Create advanced search index with TF-IDF"""
        try:
            if not self.chunks:
                self.logger.error("❌ No chunks to index")
                return False
            
            # Prepare texts for vectorization
            texts = [chunk.content for chunk in self.chunks]
            
            # Enhanced TF-IDF with better parameters
            self.vectorizer = TfidfVectorizer(
                max_features=5000,
                stop_words='english',
                ngram_range=(1, 3),  # Include bigrams and trigrams
                min_df=2,
                max_df=0.8,
                sublinear_tf=True
            )
            
            self.tfidf_matrix = self.vectorizer.fit_transform(texts)
            self.logger.info(f"✅ Search index created: {self.tfidf_matrix.shape}")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Failed to create search index: {e}")
            return False
            
    def find_relevant_chunks(self, query: str, top_k: int = 5) -> List[Tuple[DocumentChunk, float]]:
        """Find most relevant chunks with confidence scores"""
        try:
            if not self.vectorizer or self.tfidf_matrix is None:
                return []
            
            # Process query
            query_vector = self.vectorizer.transform([query])
            
            # Calculate similarities
            similarities = cosine_similarity(query_vector, self.tfidf_matrix).flatten()
            
            # Get top matches
            top_indices = similarities.argsort()[-top_k:][::-1]
            
            results = []
            for idx in top_indices:
                if similarities[idx] > 0.1:  # Minimum relevance threshold
                    chunk = self.chunks[idx]
                    chunk.confidence_score = similarities[idx]
                    results.append((chunk, similarities[idx]))
            
            return results
            
        except Exception as e:
            self.logger.error(f"❌ Search failed: {e}")
            return []
            
    def generate_comprehensive_answer(self, query: str, relevant_chunks: List[Tuple[DocumentChunk, float]]) -> Dict[str, Any]:
        """Generate comprehensive answer with proper references"""
        if not relevant_chunks:
            return {
                "answer": "I don't have enough information to answer this question based on the available documents.",
                "sources": [],
                "confidence": 0.0,
                "references": []
            }
        
        # Group chunks by source
        source_groups = defaultdict(list)
        for chunk, score in relevant_chunks:
            source_groups[chunk.source].append((chunk, score))
        
        # Build comprehensive answer
        answer_parts = []
        references = []
        sources = set()
        
        for source, chunks in source_groups.items():
            # Sort chunks by confidence
            chunks.sort(key=lambda x: x[1], reverse=True)
            
            # Take best chunk from this source
            best_chunk, best_score = chunks[0]
            
            # Format content with source reference
            content = best_chunk.content
            if best_chunk.section:
                reference = f"**{best_chunk.section}** (from {source}"
                if best_chunk.page_number:
                    reference += f", Page {best_chunk.page_number}"
                reference += ")"
            else:
                reference = f"From {source}"
                if best_chunk.page_number:
                    reference += f" (Page {best_chunk.page_number})"
            
            answer_parts.append(f"{reference}:\n{content}")
            references.append({
                "source": source,
                "section": best_chunk.section,
                "page": best_chunk.page_number,
                "confidence": best_score
            })
            sources.add(source)
        
        # Calculate overall confidence
        avg_confidence = sum(score for _, score in relevant_chunks) / len(relevant_chunks)
        
        answer = "\n\n".join(answer_parts)
        
        return {
            "answer": answer,
            "sources": list(sources),
            "confidence": avg_confidence,
            "references": references
        }
        
    def query(self, question: str) -> Dict[str, Any]:
        """Process query and return comprehensive answer"""
        try:
            self.logger.info(f"🔍 Processing query: {question}")
            
            # Find relevant chunks
            relevant_chunks = self.find_relevant_chunks(question, top_k=7)
            
            if not relevant_chunks:
                return {
                    "answer": "I couldn't find relevant information to answer this question in the available documents.",
                    "sources": [],
                    "confidence": 0.0
                }
            
            # Generate comprehensive answer
            result = self.generate_comprehensive_answer(question, relevant_chunks)
            
            self.logger.info(f"✅ Answer generated with {len(result['sources'])} sources")
            return result
            
        except Exception as e:
            self.logger.error(f"❌ Query processing failed: {e}")
            return {
                "answer": "An error occurred while processing your question.",
                "sources": [],
                "confidence": 0.0
            }
            
    def initialize(self) -> bool:
        """Initialize the premium RAG pipeline"""
        try:
            self.logger.info("🚀 Initializing Premium RAG Pipeline...")
            
            # Initialize Firebase
            self.initialize_firebase()
            
            # Download documents
            file_paths = self.download_documents_from_firebase()
            if not file_paths:
                self.logger.error("❌ No documents found")
                return False
            
            # Process documents
            if not self.process_documents(file_paths):
                self.logger.error("❌ Document processing failed")
                return False
            
            # Create search index
            if not self.create_search_index():
                self.logger.error("❌ Search index creation failed")
                return False
            
            self.logger.info("🎉 Premium RAG Pipeline initialized successfully!")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Initialization failed: {e}")
            return False
